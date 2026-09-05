#!/usr/bin/env python3
"""
ABOUTME: DOCX 学术论文后处理 — 修复 pandoc 生成的 DOCX 结构（标题居中、分页、表格宽度）
ABOUTME: 可选依赖 python-docx；未安装时降级为 no-op（记录 WARN 后返回成功）

从 opendraft-master/engine/utils/docx_post_processor.py 移植，适配 MathModelSkills 交付链路。
仅当 pandoc 路径产出 DOCX 后调用，对结构化降级版 DOCX 不生效。

用法：
    python core/tools/docx_post_processor.py paper/main.docx
    python core/tools/docx_post_processor.py paper/main.docx --institution "XX大学"
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Any, Dict, Optional


# ---------------------------------------------------------------------------
# 可选依赖探测
# ---------------------------------------------------------------------------
_DOCX_AVAILABLE = False
try:
    from docx import Document  # noqa: F401
    from docx.shared import Pt, Inches  # noqa: F401
    from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH  # noqa: F401
    _DOCX_AVAILABLE = True
except ImportError:
    pass


# ---------------------------------------------------------------------------
# 公开 API
# ---------------------------------------------------------------------------

def insert_academic_structure(
    docx_path: Path,
    verbose: bool = False,
    options: Optional[Dict[str, Any]] = None,
) -> bool:
    """对 pandoc 生成的 DOCX 插入学术论文结构。"""
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX 文件不存在: {docx_path}")

    if not _DOCX_AVAILABLE:
        if verbose:
            print("[docx_post_processor] python-docx 未安装，跳过后处理（不影响 DOCX 交付）")
        return True

    try:
        return _post_process(docx_path, verbose, options or {})
    except Exception as e:
        if verbose:
            print(f"[docx_post_processor] 后处理失败: {e}")
            import traceback
            traceback.print_exc()
        return False


# ---------------------------------------------------------------------------
# 内部实现（依赖 python-docx）
# ---------------------------------------------------------------------------

def _post_process(docx_path: Path, verbose: bool, options: Dict[str, Any]) -> bool:
    from docx import Document
    from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH

    doc = Document(docx_path)

    title_idx, date_idx = _find_title_block(doc)
    if title_idx is None:
        if verbose:
            print("[docx_post_processor] 未找到 Title 样式，跳过后处理")
        return True

    if options.get("institution"):
        _insert_institution_block(doc, title_idx, options, verbose)
        title_idx, date_idx = _find_title_block(doc)

    _center_title_block(doc, title_idx, date_idx)

    if options:
        _insert_metadata_after_date(doc, date_idx, options, verbose)
        _, date_idx = _find_title_block(doc)

    cover_end_idx = _find_cover_end(doc)
    if cover_end_idx is not None:
        _insert_page_break_after(doc, cover_end_idx)

    abstract_end_idx = _find_abstract_end(doc)
    if abstract_end_idx is not None:
        _insert_page_break_after(doc, abstract_end_idx)

    _fix_table_widths(doc, verbose)
    doc.save(docx_path)
    return True


def _find_title_block(doc) -> tuple:
    title_idx = None
    date_idx = None
    for i, para in enumerate(doc.paragraphs[:20]):
        style = para.style.name if para.style else ""
        if style == "Title" and title_idx is None:
            title_idx = i
        elif style == "Date":
            date_idx = i
            break
    return title_idx, date_idx


def _center_title_block(doc, title_idx, date_idx) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if title_idx is None or date_idx is None:
        return
    for i in range(title_idx, date_idx + 1):
        if i < len(doc.paragraphs):
            doc.paragraphs[i].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _find_cover_end(doc):
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip().lower()
        if "abstract" in text and "Heading" in style:
            return i - 1
        if style.startswith("Heading") and i > 5:
            return i - 1
    return None


def _find_abstract_end(doc):
    in_abstract = False
    last_abstract_para = None
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if "abstract" in text.lower() and "Heading" in style:
            in_abstract = True
            continue
        if in_abstract:
            if style == "Heading 1" and text and (text[0].isdigit() or text.startswith("1")):
                return last_abstract_para
            last_abstract_para = i
    return None


def _insert_institution_block(doc, title_idx: int, options: Dict, verbose: bool) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    items = []
    if options.get("department"):
        items.append(("department", options["department"], True, 11))
    if options.get("faculty"):
        items.append(("faculty", options["faculty"], False, 11))
    if options.get("institution"):
        items.append(("institution", options["institution"].upper(), False, 14))
    insert_count = 0
    for _name, text, italic, size in items:
        para = doc.paragraphs[title_idx]
        new_para = para.insert_paragraph_before(text)
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in new_para.runs:
            run.font.size = Pt(size)
            if italic:
                run.font.italic = True
            if _name == "institution":
                run.font.small_caps = True
        insert_count += 1
    if insert_count > 0:
        para = doc.paragraphs[title_idx + insert_count]
        para.insert_paragraph_before("")
    if verbose and insert_count > 0:
        print(f"[docx_post_processor] 已插入机构信息（{insert_count} 行）")


def _insert_metadata_after_date(doc, date_idx: int, options: Dict, verbose: bool) -> None:
    if date_idx is None:
        return
    additions = 0
    date_para = doc.paragraphs[date_idx]
    insert_after = date_para
    if options.get("project_type"):
        insert_after = _insert_para_after(insert_after, "")
        insert_after = _insert_para_after(insert_after, options["project_type"].upper(), size=11, small_caps=True)
        additions += 2
    if options.get("instructor"):
        insert_after = _insert_para_after(insert_after, "")
        insert_after = _insert_para_after(insert_after, f"First Supervisor: {options['instructor']}", size=10)
        additions += 2
    if options.get("second_examiner"):
        insert_after = _insert_para_after(insert_after, f"Second Examiner: {options['second_examiner']}", size=10)
        additions += 1
    if options.get("location"):
        insert_after = _insert_para_after(insert_after, "")
        insert_after = _insert_para_after(insert_after, options["location"], size=10)
        additions += 2
    if verbose and additions > 0:
        print(f"[docx_post_processor] 已插入附加元数据（{additions} 行）")


def _insert_para_after(after_para, text: str, size: int = 11, bold: bool = False,
                       italic: bool = False, small_caps: bool = False):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    parent = after_para._element.getparent()
    index = list(parent).index(after_para._element)
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    new_p.append(pPr)
    if text:
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(size * 2))
        rPr.append(szCs)
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))
        if small_caps:
            rPr.append(OxmlElement("w:smallCaps"))
        run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        new_p.append(run)
    parent.insert(index + 1, new_p)
    doc = after_para._element.getparent().getparent()
    for para in after_para._element.getparent().iterchildren(qn("w:p")):
        if para is new_p:
            from docx.text.paragraph import Paragraph
            return Paragraph(new_p, after_para._parent)
    return after_para


def _insert_page_break_after(doc, para_index: int) -> None:
    from docx.enum.text import WD_BREAK
    if para_index is None or para_index >= len(doc.paragraphs):
        return
    target_para = doc.paragraphs[para_index]
    run = target_para.add_run()
    run.add_break(WD_BREAK.PAGE)


def _fix_table_widths(doc, verbose: bool = False) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt
    tables_fixed = 0
    for table in doc.tables:
        num_cols = len(table.columns)
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")
        tblLayout = tblPr.find(qn("w:tblLayout"))
        if tblLayout is None:
            tblLayout = OxmlElement("w:tblLayout")
            tblPr.append(tblLayout)
        tblLayout.set(qn("w:type"), "autofit")
        table.autofit = True
        if num_cols >= 5:
            font_size = Pt(8)
        elif num_cols >= 4:
            font_size = Pt(9)
        else:
            font_size = Pt(10)
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tcPr = tc.tcPr
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    tc.insert(0, tcPr)
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is not None:
                    tcPr.remove(tcW)
                if col_idx == 0:
                    noWrap = tcPr.find(qn("w:noWrap"))
                    if noWrap is None:
                        noWrap = OxmlElement("w:noWrap")
                        tcPr.append(noWrap)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = font_size
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after = Pt(2)
        tables_fixed += 1
    if verbose and tables_fixed > 0:
        print(f"[docx_post_processor] 已修复 {tables_fixed} 个表格（自动宽度，{font_size.pt}pt 字体）")


# ---------------------------------------------------------------------------
# CLI 入口
# ---------------------------------------------------------------------------

def main() -> int:
    ap = argparse.ArgumentParser(description="DOCX 学术论文后处理（修复标题居中/分页/表格宽度）")
    ap.add_argument("docx", help="DOCX 文件路径")
    ap.add_argument("--institution", help="机构名（大写、居中置于标题前）")
    ap.add_argument("--faculty", help="学院名")
    ap.add_argument("--department", help="系名（斜体）")
    ap.add_argument("--instructor", help="指导教师")
    ap.add_argument("--verbose", "-v", action="store_true")
    args = ap.parse_args()

    docx_path = Path(args.docx)
    if not docx_path.exists():
        print(f"[docx_post_processor] 文件不存在: {docx_path}", file=sys.stderr)
        return 1

    options = {k: v for k, v in vars(args).items()
               if k not in ("docx", "verbose") and v is not None}
    ok = insert_academic_structure(docx_path, verbose=args.verbose, options=options)
    return 0 if ok else 1


if __name__ == "__main__":
    sys.exit(main())
